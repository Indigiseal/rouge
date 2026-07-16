from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Users\tayak\Downloads\event-stories.md")
OUTPUT = Path(r"C:\Users\tayak\Downloads\Event Stories - Easy Reading.docx")


FONT = "Calibri"
INK = RGBColor(31, 31, 31)
MUTED = RGBColor(92, 92, 92)
ACCENT = RGBColor(46, 116, 181)
CHOICE = RGBColor(112, 48, 160)


def set_run_font(run, name=FONT, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.3

    for name, size, color, before, after in [
        ("Heading 1", 19, ACCENT, 18, 7),
        ("Heading 2", 15, ACCENT, 12, 5),
        ("Heading 3", 13, CHOICE, 10, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = name != "Heading 1"
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=2, line=1.0)
    r = p.add_run("Event Stories")
    set_run_font(r, size=28, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=18, line=1.15)
    r = p.add_run("A cleaner, easier-to-read version")
    set_run_font(r, size=12, italic=True, color=MUTED)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=16, line=1.3)
    r = p.add_run(
        "Each section is one dungeon event. The choices are separated so the story can be read a little at a time."
    )
    set_run_font(r, size=12, color=INK)


def add_choice_block(doc, label, text_color=CHOICE):
    p = doc.add_paragraph()
    set_paragraph_shading(p, "F7F2FB")
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    set_paragraph_spacing(p, before=7, after=7, line=1.15)
    r = p.add_run(label)
    set_run_font(r, size=12, bold=True, color=text_color)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=8, line=1.3)
    r = p.add_run(text)
    set_run_font(r, size=12, color=INK)


def normalize_heading(text):
    text = text.strip()
    prefixes = [
        "Choice:",
        "Drag Outcome:",
        "Random Outcome:",
        "Alternate fatal version:",
        "Alternate version:",
        "Success:",
        "Failure:",
    ]
    for prefix in prefixes:
        if text.startswith(prefix):
            return text
    return text


def build():
    md = SOURCE.read_text(encoding="utf-8-sig")
    doc = Document()
    style_document(doc)
    add_title(doc)

    first_story = True
    skip_intro_note = False

    for raw_line in md.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line == "# Event Stories":
            skip_intro_note = True
            continue
        if skip_intro_note and line.startswith("Narrative-only extraction"):
            skip_intro_note = False
            continue

        heading_match = re.match(r"^(#{2,5})\s+(.*)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = normalize_heading(heading_match.group(2))

            if level == 2:
                first_story = False
                p = doc.add_paragraph(text, style="Heading 1")
                p.paragraph_format.page_break_before = False
            elif text.startswith(("Choice:", "Drag Outcome:", "Random Outcome:")):
                add_choice_block(doc, text)
            elif text in ("Success:", "Failure:", "Alternate fatal version:", "Alternate version:"):
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=6, after=4, line=1.15)
                r = p.add_run(text)
                set_run_font(r, size=12, bold=True, italic=True, color=MUTED)
            else:
                p = doc.add_paragraph(text, style="Heading 2" if level == 3 else "Heading 3")
            continue

        add_body_paragraph(doc, line)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer, before=0, after=0, line=1.0)
    r = footer.add_run("Event Stories")
    set_run_font(r, size=9, color=MUTED)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)
